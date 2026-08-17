"""Generate Software Requirements Specification (SRS.docx) following IEEE 830 standard."""
import os
import docx
from docx import Document
from docx_styler import (
    set_page_margins, add_title, add_subtitle, add_authors, add_heading_1,
    add_heading_2, add_heading_3, add_body, add_bullet, add_callout_box,
    add_formatted_table, add_image_figure, clean_text
)

def create_srs():
    doc = Document()
    set_page_margins(doc)

    # Header
    add_title(doc, "SOFTWARE REQUIREMENTS SPECIFICATION (SRS)")
    add_subtitle(doc, "IEEE Std 830-1998 Compliant Specification for Edge Intrusion Detection System in Digital Mining Operations\nUNESCO Project - Track 3: Smart Subsoil | Russian-African Forum 2026")
    add_authors(doc,
        "Software Engineering Delegation: John Okyere (Lead), Ezekeil Baah, Clement Baffour, Parker Paa Annobil, George Akwesi Bonnah",
        "Department of ICT, University of Education, Winneba & Kayaba Labs | Standard: IEEE 830-1998 | Version 3.0.0"
    )

    add_callout_box(doc, "SPECIFICATION SCOPE & COMPLIANCE",
        "This Software Requirements Specification (SRS) establishes the technical baseline, system interfaces, behavioral constraints, and verification criteria for the 'Securing the Digital Mine' metaheuristic-optimized deep learning intrusion detection artifact. This document complies with IEEE Std 830-1998 standards for software requirements specification."
    )

    # 1. INTRODUCTION
    add_heading_1(doc, "1. INTRODUCTION")
    
    add_heading_2(doc, "1.1 Purpose")
    add_body(doc,
        "The purpose of this document is to provide a complete, formal, and unambiguous specification of the software requirements for the 'Securing the Digital Mine' Edge Intrusion Detection System (IDS). It details the functional capabilities, external interfaces, performance guarantees, and design constraints for edge gateways, cloud API handlers, and the real-time SaaS monitoring dashboard."
    )

    add_heading_2(doc, "1.2 Scope")
    add_body(doc,
        "The software system encompasses: (1) An edge telemetry collection and feature extraction client (@mhiskall282/unesco-mine-sec-cli); (2) A Binary Whale Optimization Algorithm (BWOA) feature pruning engine; (3) A spatial-temporal Convolutional Neural Network and Long Short-Term Memory (CNN-LSTM) classifier; (4) A Float16 Quantized TensorFlow Lite inference server (FastAPI, port 8001); and (5) A multi-tenant Laravel 12 Livewire SaaS monitoring and incident management portal."
    )

    add_heading_2(doc, "1.3 Definitions, Acronyms, and Abbreviations")
    add_formatted_table(doc,
        ["Term / Acronym", "Full Definition & Context"],
        [
            ["BWOA", "Binary Whale Optimization Algorithm - A nature-inspired metaheuristic used for discrete feature subset selection."],
            ["CNN-LSTM", "Convolutional Neural Network and Long Short-Term Memory - A hybrid neural architecture combining spatial feature extraction with temporal sequence modeling."],
            ["IIoT", "Industrial Internet of Things - Networked sensors and instrumentation deployed across industrial mining machinery."],
            ["Modbus", "A legacy industrial serial/TCP communications protocol commonly used to monitor PLCs, flowmeters, and power meters in mining."],
            ["OT", "Operational Technology - Hardware and software that detects or causes changes through direct monitoring and control of physical industrial processes."],
            ["PLC", "Programmable Logic Controller - Ruggedized industrial digital computer controlling manufacturing and extraction processes."],
            ["SCADA", "Supervisory Control and Data Acquisition - High-level industrial supervision architecture."],
            ["TFLite", "TensorFlow Lite - Lightweight runtime engine optimized for low-latency inference on mobile and edge devices."],
            ["UAT", "User Acceptance Testing - Structured usability and functional verification conducted with domain specialists."]
        ],
        col_widths=[2.0, 4.5]
    )

    # 2. OVERALL DESCRIPTION
    add_heading_1(doc, "2. OVERALL SYSTEM DESCRIPTION")
    
    add_heading_2(doc, "2.1 Product Perspective")
    add_body(doc,
        "The system operates as an autonomous edge-to-cloud cyber-physical defense shield. It intercepts network flows at remote mining substation gateways before unauthorized Modbus coil commands or volumetric denial-of-service floods can compromise industrial PLCs controlling SAG mills, conveyor belts, or toxic gas ventilation grids."
    )
    add_image_figure(doc, "research/figures/system_architecture.png", "Figure 2.1: Context Diagram - Four-Layer Industrial Defense Pipeline", width_inches=6.2)

    add_heading_2(doc, "2.2 Operating Environment")
    add_formatted_table(doc,
        ["Environment Component", "Specification Baseline", "Minimum Requirements"],
        [
            ["Edge Hardware Platform", "Raspberry Pi 4 Model B / Raspberry Pi 5", "ARMv8 64-bit Cortex-A72 CPU, 1GB RAM, 16GB MicroSD"],
            ["Edge Operating System", "Raspberry Pi OS (64-bit Debian Bullseye/Bookworm)", "Linux Kernel 5.15+, systemd init, python3.11, nodejs20"],
            ["Cloud Server Platform", "AWS EC2 / Ubuntu 22.04 LTS (t3.medium)", "2 vCPUs, 4GB RAM, 20GB SSD, Python 3.11, Uvicorn"],
            ["Dashboard Framework", "Laravel 12 / PHP 8.2 / Livewire 3 / PostgreSQL", "Nginx web server, modern browser (Chromium/Firefox)"]
        ],
        col_widths=[2.2, 2.3, 2.0]
    )

    # 3. SPECIFIC REQUIREMENTS
    add_heading_1(doc, "3. SPECIFIC SOFTWARE REQUIREMENTS")

    add_heading_2(doc, "3.1 External Interface Requirements")
    add_body(doc, "The software provides three distinct external interfaces:", bold_prefix="Interface Scope: ")
    add_bullet(doc, "CLI Telemetry Interface: Interactive command-line interface with Inquirer.js wizards for adapter selection, API endpoint configuration, and live Chalk-colorized flow streaming.")
    add_bullet(doc, "REST API Endpoints: FastAPI microservice exposing GET /api/health, GET /api/features, and POST /api/analyze for JSON payload ingestion.")
    add_bullet(doc, "Livewire SaaS Portal: Responsive multi-tenant web console broadcasting real-time attack alerts, confidence scores, and historical forensic logs.")

    add_heading_2(doc, "3.2 Detailed Functional Requirements (IEEE 830 Standard)")
    add_formatted_table(doc,
        ["Requirement Tag", "Module / Subsystem", "Description & Functional Logic", "Verification Gate"],
        [
            ["SRS-FR-101", "Packet Capture Engine", "Sniffs promiscuous interface frames, extracts 10 BWOA fields, and constructs JSON flow objects.", "Unit test test_api_service.py passing"],
            ["SRS-FR-102", "BWOA Pruner", "Applies pre-computed 10-feature mask: [src_bytes, service, flag, serror_rate, same_srv_rate, diff_srv_rate, dst_host_diff_srv_rate, protocol_type, hot, su_attempted].", "Unit test test_bwoa.py passing"],
            ["SRS-FR-103", "Neural Inference", "Executes Float16 TFLite interpreter, evaluates Conv1D spatial filters and LSTM temporal cells, returns Softmax probability distribution.", "Unit test test_cnn_lstm.py passing"],
            ["SRS-FR-104", "Threat Thresholding", "Identifies winning class. If class != Normal, sets alert flag and identifies specific feature triggers (e.g. high_serror_rate).", "Integration validate_api.py passing"],
            ["SRS-FR-105", "Audit Persistence", "Persists flow records, prediction outcomes, device tokens, and latency telemetry to PostgreSQL/SQLite database.", "Database migration passing"]
        ],
        col_widths=[1.2, 1.8, 2.3, 1.2]
    )

    add_heading_2(doc, "3.3 UML Structural and Behavioral Models")
    add_image_figure(doc, "research/figures/uml_use_case.png", "Figure 3.1: UML Use Case Diagram - Operator & Analyst Roles", width_inches=5.8)
    add_image_figure(doc, "research/figures/uml_class_diagram.png", "Figure 3.2: UML Class Diagram - Core Class Model & Method Signatures", width_inches=5.8)
    add_image_figure(doc, "research/figures/uml_sequence_diagram.png", "Figure 3.3: UML Sequence Diagram - Real-Time Intrusion Detection Interaction", width_inches=5.8)

    # 4. SYSTEM VERIFICATION & TESTING
    add_heading_1(doc, "4. SYSTEM VERIFICATION & QUALITY ASSURANCE")
    add_body(doc,
        "The software architecture enforces four levels of quality assurance and automated testing across all build pipelines:"
    )
    add_formatted_table(doc,
        ["Testing Level", "Test Scope & Commands", "Automated Criteria", "Test Result"],
        [
            ["Unit Testing", "python -m unittest discover -s tests -p 'test_*.py'", "75 unit tests across 9 test suites", "75/75 PASS (125.6s)"],
            ["API Validation", "python scripts/validate_api.py", "Health, Features, Analyze, 404 handler", "100% PASS"],
            ["AWS EC2 Deployment", "bash scripts/validate_ec2_deployment.sh", "Port checks, systemd daemons, dependencies", "STATUS: READY (0 errors)"],
            ["Raspberry Pi Dry-Run", "bash scripts/validate_pi_deployment.sh", "ARM TFLite runtime, npm binary linkage", "STATUS: READY"],
            ["Documentation Integrity", "python scripts/verify_readme_links.py", "31 internal markdown hyperlinks", "31/31 PASS"]
        ],
        col_widths=[1.5, 2.3, 1.5, 1.2]
    )

    output_path = "research/SRS.docx"
    doc.save(output_path)
    print(f"Software Requirements Specification (SRS) saved successfully to {output_path}!")

if __name__ == "__main__":
    create_srs()

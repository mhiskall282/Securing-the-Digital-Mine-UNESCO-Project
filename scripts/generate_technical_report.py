"""Generate the comprehensive Technical Report in docx format with Appendices A-E."""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_technical_report():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 29, 58)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(71, 85, 105)

    def add_authors():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run("Engineering Delegation: John Okyere, Ezekeil Baah, Clement Baffour, Parker Paa Annobil, George Akwesi Bonnah\nUniversity of Education, Winneba & Kayaba Labs | UNESCO Russian-African Forum 2026")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 82, 155)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 82, 155)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Times New Roman'
            run_b.font.size = Pt(10.5)
            run_b.font.bold = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Times New Roman'
            run_b.font.size = Pt(10)
            run_b.font.bold = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        return p

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        run_t = p.add_run(title + "\n")
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10)
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(0, 82, 155)
        
        run_b = p.add_run(text)
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(9.5)
        run_b.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "0F172A") # Dark slate
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(226, 232, 240)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_image_box(image_path, caption, width_inches=5.8):
        if os.path.exists(image_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(image_path, width=Inches(width_inches))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            run_c = p_cap.add_run(caption)
            run_c.font.name = 'Times New Roman'
            run_c.font.size = Pt(9)
            run_c.font.italic = True
            run_c.font.color.rgb = RGBColor(71, 85, 105)

    def add_styled_table(headers, rows, col_widths=None):
        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Header Row
        hdr_cells = tbl.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "00529B")
            set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data Rows
        for r_idx, row_data in enumerate(rows):
            row_cells = tbl.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = str(val)
                set_cell_background(row_cells[c_idx], bg_color)
                set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
                p = row_cells[c_idx].paragraphs[0]
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif c_idx == len(row_data) - 1 and ("PASS" in str(val) or "Yes" in str(val)):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9)
                    if "PASS" in str(val):
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(16, 185, 129)
        
        if col_widths:
            for row in tbl.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # DOCUMENT CONTENT
    # -------------------------------------------------------------
    add_title("TECHNICAL REPORT & DEPLOYMENT SPECIFICATION")
    add_subtitle("Securing the Digital Mine: Edge-Ready Metaheuristic Optimized Deep Learning Intrusion Detection")
    add_authors()

    # SECTION 1: EXECUTIVE SUMMARY
    add_heading_1("1. EXECUTIVE SUMMARY")
    add_body(
        "Modern mineral resource operations across Africa and the Russian Federation are undergoing rapid digital transformation under 'Smart Subsoil' initiatives. Heavy mining equipment—including ball mills, cone crushers, autonomous haul trucks, and tailings pumps—are increasingly connected to operational technology (OT) SCADA networks. While this connectivity drives massive productivity and safety gains, it creates severe vulnerabilities to malicious cyber disruption. Unplanned industrial downtime in mining facilities costs between USD $50,000 and $500,000 per hour, while cyber-physical attacks on ventilation or tailings management represent direct threats to human life."
    )
    add_body(
        "Commercial enterprise cybersecurity solutions cannot solve this problem because they rely on static signatures, cloud-based inspection requiring high-bandwidth links, and heavy computation that exceeds the 100-millisecond control loop deadlines of industrial PLCs. Remote African mining operations operating on solar power and satellite links require lightweight, offline-capable, real-time edge security."
    )
    add_body(
        "This technical report details the architecture, implementation, and deployment of a three-tier intrusion detection framework combining Binary Whale Optimization (BWOA) with a spatial-temporal CNN-LSTM neural network. By pruning network telemetry from 41 features to 10 (75.61% reduction) and applying Float16 quantization, the system achieves an inference latency of 0.76 milliseconds on a standard Raspberry Pi 4 edge device while maintaining 70.56% multi-class accuracy and 96.89% precision on benign flows. The complete open-source solution enables mining operators to achieve immediate cyber resilience at negligible capital cost."
    )

    # SECTION 2: TECHNICAL ARCHITECTURE
    add_heading_1("2. TECHNICAL ARCHITECTURE & DEEP DIVE")
    add_image_box("research/figures/system_architecture.png", "Figure 2.1: Four-Layer End-to-End System Architecture", width_inches=6.0)

    add_heading_2("2.1 Layer 1: Edge Telemetry Ingestion Agent")
    add_body(
        "The edge telemetry ingestion agent (@mhiskall282/unesco-mine-sec-cli) is packaged as a high-performance Node.js service designed to run directly on edge gateways, micro-PLCs, or local aggregation servers. It captures raw network frames from promiscuous network adapters (Ethernet, Wi-Fi, TAP bridges), parses IP/TCP/UDP headers, and streams structured JSON payloads containing the 10 BWOA-selected features over HTTP/HTTPS."
    )

    add_heading_2("2.2 Layer 2: BWOA Feature Selection Engine")
    add_body(
        "The Binary Whale Optimization Algorithm (BWOA) conducts metaheuristic exploration of the 2^41 discrete feature space. Continuous velocity vectors are mapped to bit-flip probabilities using a V-shaped transfer function with a strict 75% accuracy floor constraint to ensure high discrimination on minority attack classes."
    )
    add_image_box("research/figures/bwoa_convergence.png", "Figure 2.2: BWOA Optimization Convergence across 100 Iterations", width_inches=5.4)

    add_heading_2("2.3 Layer 3: Spatial-Temporal CNN-LSTM Classifier")
    add_body(
        "The deep learning engine combines a 1D Convolutional Neural Network (Conv1D) layer (64 filters, kernel size 3) with a Long Short-Term Memory (LSTM) layer (256 units). The Conv1D filters isolate local spatial relationships between byte values, while the LSTM recurrent units track temporal state transitions across sequential polling cycles."
    )
    add_image_box("research/figures/cnn_lstm_architecture.png", "Figure 2.3: CNN-LSTM Neural Network Architecture Flowchart", width_inches=5.8)

    add_heading_2("2.4 Layer 4: Model Server & SaaS Livewire Dashboard")
    add_body(
        "The inference microservice runs an asynchronous FastAPI server (port 8001) executing the quantized Float16 TFLite model. Threat predictions are ingested by a multi-tenant Laravel 12 Livewire SaaS dashboard that broadcasts live alerts, maps device telemetry, and registers incident logs."
    )

    # SECTION 3: STEP-BY-STEP DEPLOYMENT GUIDE
    add_heading_1("3. DEPLOYMENT & INSTALLATION GUIDE")

    add_heading_2("3.1 Edge Gateway Deployment (Raspberry Pi 4B / Pi 5)")
    add_body("Follow these steps to deploy the complete intrusion detection stack on a Linux edge gateway:")
    add_code_block(
        "# 1. Clone the project repository\n"
        "git clone https://github.com/mhiskall282/Securing-the-Digital-Mine-UNESCO-Project.git /opt/unesco-project\n"
        "cd /opt/unesco-project\n\n"
        "# 2. Run the automated deployment installer\n"
        "chmod +x scripts/deploy_raspberry_pi.sh\n"
        "./scripts/deploy_raspberry_pi.sh\n\n"
        "# 3. Configure and launch the global CLI agent\n"
        "npm config set @mhiskall282:registry https://npm.pkg.github.com\n"
        "npm install -g @mhiskall282/unesco-mine-sec-cli\n"
        "unesco-mine-sec-cli --url http://127.0.0.1:8001/api/analyze --interface eth0"
    )

    add_heading_2("3.2 Cloud Inference Server Deployment (AWS EC2 / Ubuntu 22.04)")
    add_code_block(
        "# 1. Provision Ubuntu 22.04 LTS instance (t3.medium recommended)\n"
        "# 2. Clone repository and run automated AWS setup\n"
        "git clone https://github.com/mhiskall282/Securing-the-Digital-Mine-UNESCO-Project.git /opt/unesco-project\n"
        "cd /opt/unesco-project\n"
        "chmod +x scripts/deploy_aws_ec2.sh\n"
        "./scripts/deploy_aws_ec2.sh\n\n"
        "# 3. Verify service status and endpoints\n"
        "curl http://localhost:8001/api/health\n"
        "curl http://localhost:8001/api/features"
    )

    # SECTION 4: EVALUATION RESULTS
    add_heading_1("4. COMPREHENSIVE EVALUATION RESULTS")
    add_styled_table(
        ["Metric Category", "Baseline (41 Feat)", "BWOA v3 (10 Feat)", "BWOA Float16", "SWaT Transfer"],
        [
            ["Input Dimensions", "41 Features", "10 Features", "10 Features", "51 Sensor Feat"],
            ["Test Accuracy", "77.70%", "70.56%", "70.56%", "59.95%"],
            ["Macro F1-Score", "0.7571", "0.7127", "0.7127", "0.5966"],
            ["AUC-ROC Score", "0.9359", "0.8471", "0.8471", "0.8650"],
            ["Inference Latency", "157.66 ms", "35.60 ms", "0.76 ms", "0.12 ms"],
            ["Model Size", "1.86 MB", "4.88 MB", "0.82 MB", "1.76 MB"],
            ["SCADA Compliance", "FAIL (>100ms)", "PASS (<100ms)", "PASS (<100ms)", "PASS (<100ms)"]
        ],
        col_widths=[1.8, 1.3, 1.3, 1.3, 1.3]
    )
    add_image_box("research/figures/confusion_matrix.png", "Figure 4.1: Confusion Matrix on NSL-KDD Held-Out Test Set (22,544 Samples)", width_inches=5.4)
    add_image_box("research/figures/latency_comparison_barchart.png", "Figure 4.2: Latency Profile Comparison across IDS Implementations", width_inches=5.6)

    # APPENDICES
    add_heading_1("APPENDICES")

    # Appendix A
    add_heading_2("APPENDIX A: BWOA Optimization Pseudocode")
    add_code_block(
        "Algorithm 1: Binary Whale Optimization Algorithm (BWOA) with Accuracy Floor\n"
        "----------------------------------------------------------------------------\n"
        "Input : Feature matrix X in R^{N x D}, labels y in {0, ..., C-1}\n"
        "        Parameters: n_agents=30, max_iter=100, alpha=0.3, min_acc=0.75\n"
        "Output: Best binary feature mask X_best in {0, 1}^D\n\n"
        "1. Initialize population of whale positions X_i in {0, 1}^D randomly for i = 1 to n_agents\n"
        "2. Evaluate fitness for each agent: Fit_i = alpha * (1 - Acc_i) + (1 - alpha) * (|X_i|/D) + Penalty_i\n"
        "3. Identify leader whale X_best with minimum fitness score\n"
        "4. while t < max_iter do:\n"
        "5.     a = 2 - 2 * (t / max_iter)  // Linear parameter decay\n"
        "6.     for each agent i do:\n"
        "7.         p = rand(), r1 = rand(), r2 = rand(), l = rand(-1, 1)\n"
        "8.         A = 2 * a * r1 - a,  C = 2 * r2\n"
        "9.         if p < 0.5 then:\n"
        "10.            if |A| < 1 then:\n"
        "11.                D_vec = |C * X_best - X_i|\n"
        "12.                X_cont = X_best - A * D_vec  // Shrinking encircling\n"
        "13.            else:\n"
        "14.                X_rand = select_random_whale()\n"
        "15.                D_vec = |C * X_rand - X_i|\n"
        "16.                X_cont = X_rand - A * D_vec  // Exploration\n"
        "17.        else:\n"
        "18.            D_prime = |X_best - X_i|\n"
        "19.            X_cont = D_prime * exp(b * l) * cos(2 * pi * l) + X_best  // Spiral search\n"
        "20.        // Apply V-shaped binary transfer function\n"
        "21.        for each dimension d do:\n"
        "22.            V_val = |X_cont[d] / sqrt(1 + X_cont[d]^2)|\n"
        "23.            if rand() < V_val then X_i[d] = 1 - X_i[d]\n"
        "24.        Re-evaluate fitness Fit_i and update X_best\n"
        "25.    t = t + 1\n"
        "26. return X_best"
    )

    # Appendix B
    add_heading_2("APPENDIX B: CNN-LSTM Hyperparameters & Architecture Specifications")
    add_styled_table(
        ["Layer Type", "Output Shape", "Param Count", "Activation", "Regularization / Details"],
        [
            ["Input Layer", "(None, 10, 1)", "0", "-", "10 BWOA Features reshaped for 1D convolution"],
            ["Conv1D", "(None, 10, 64)", "256", "ReLU", "Filters=64, Kernel Size=3, Padding='same'"],
            ["BatchNormalization", "(None, 10, 64)", "256", "-", "Normalizes activations across mini-batches"],
            ["Spatial Dropout", "(None, 10, 64)", "0", "-", "Dropout Rate = 0.3 to prevent overfitting"],
            ["LSTM Layer", "(None, 256)", "328,704", "Tanh", "Units=256, Recurrent Activation='sigmoid'"],
            ["Dense (Hidden)", "(None, 64)", "16,448", "ReLU", "Fully connected representation layer"],
            ["Dropout", "(None, 64)", "0", "-", "Dropout Rate = 0.2"],
            ["Dense (Output)", "(None, 5)", "325", "Softmax", "Multi-class probabilities (5 attack classes)"]
        ],
        col_widths=[1.5, 1.2, 1.0, 1.0, 2.3]
    )

    # Appendix C
    add_heading_2("APPENDIX C: Complete CICFlowMeter to NSL-KDD Feature Mapping")
    add_styled_table(
        ["Idx", "NSL-KDD Feature", "BWOA Status", "CICFlowMeter Equivalent", "Description"],
        [
            ["1", "duration", "Pruned", "Flow Duration", "Connection duration in seconds"],
            ["2", "protocol_type", "SELECTED (8)", "Protocol", "Network layer protocol (TCP/UDP/ICMP)"],
            ["3", "service", "SELECTED (2)", "Dst Port / App Protocol", "Destination service (HTTP, Modbus, Private)"],
            ["4", "flag", "SELECTED (3)", "TCP Flags Count", "Connection completion status (SF, S0, REJ)"],
            ["5", "src_bytes", "SELECTED (1)", "Total Fwd Packets Bytes", "Bytes sent from source to destination"],
            ["6", "dst_bytes", "Pruned", "Total Bwd Packets Bytes", "Bytes sent from destination to source"],
            ["7", "hot", "SELECTED (9)", "Sensitive File Access", "Indicators of sensitive system access"],
            ["8", "su_attempted", "SELECTED (10)", "Privilege Escalation Flag", "Root/admin privilege escalation attempts"],
            ["9", "serror_rate", "SELECTED (4)", "SYN Error Rate", "Proportion of connections with SYN errors"],
            ["10", "same_srv_rate", "SELECTED (5)", "Same Service Ratio", "Proportion of connections to same service"],
            ["11", "diff_srv_rate", "SELECTED (6)", "Diff Service Ratio", "Proportion of connections to different services"],
            ["12", "dst_host_diff_srv_rate", "SELECTED (7)", "Dst Host Diff Srv Rate", "Destination host service dispersion"]
        ],
        col_widths=[0.5, 1.8, 1.3, 1.8, 1.8]
    )

    # Appendix D
    add_heading_2("APPENDIX D: User Acceptance Testing (UAT) Questionnaire")
    add_body(
        "Domain specialists (3 cybersecurity analysts and 2 mining OT engineers) scored the platform on a 1–5 Likert scale across five criteria:\n"
        "1. Alert Clarity: Are threat predictions understandable without deep cybersecurity expertise?\n"
        "2. Dashboard Usability: Is live telemetry streaming intuitive and actionable for control room operators?\n"
        "3. Setup Ease: Can the sniffer CLI agent be deployed on a new device in under 5 minutes?\n"
        "4. Confidence Scoring: Does the confidence percentage provide actionable insight for incident triage?\n"
        "5. Operational Fit: Would you recommend this solution for deployment on remote African mining sites?"
    )

    # Appendix E
    add_heading_2("APPENDIX E: Repository & Reproducibility Links")
    add_bullet("GitHub Main Repository: https://github.com/mhiskall282/Securing-the-Digital-Mine-UNESCO-Project", bold_prefix="Source Code: ")
    add_bullet("NPM Package (GitHub Packages): @mhiskall282/unesco-mine-sec-cli", bold_prefix="CLI Package: ")
    add_bullet("Google Colab Training Pipeline: notebooks/00_colab_setup_and_train.ipynb", bold_prefix="Google Colab: ")
    add_bullet("Live Dashboard URL: https://minesec-dashboard-prod.onrender.com", bold_prefix="Live Demo: ")

    output_path = "research/technical_report.docx"
    doc.save(output_path)
    print(f"Technical Report saved successfully to {output_path}!")

if __name__ == "__main__":
    create_technical_report()

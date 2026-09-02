"""Generate Product Requirements Document (PRD.docx) for Securing the Digital Mine IDS."""
import os
import docx
from docx import Document
from docx_styler import (
    set_page_margins, add_title, add_subtitle, add_authors, add_heading_1,
    add_heading_2, add_heading_3, add_body, add_bullet, add_callout_box,
    add_formatted_table, add_image_figure, clean_text
)

def create_prd():
    doc = Document()
    set_page_margins(doc)

    # Document Header
    add_title(doc, "PRODUCT REQUIREMENTS DOCUMENT (PRD)")
    add_subtitle(doc, "Securing the Digital Mine: Edge-Deployable OT/IIoT Intrusion Detection System\nUNESCO Project - Track 3: Smart Subsoil | Russian-African Forum 2026")
    add_authors(doc, 
        "Product Engineering Team: John Okyere (Lead), Ezekeil Baah, Clement Baffour, Parker Paa Annobil, George Akwesi Bonnah",
        "University of Education, Winneba & UEW Innovation Hub | Version 3.0.0 | Status: Approved for Deployment"
    )

    # Callout Summary
    add_callout_box(doc, "PRODUCT MISSION STATEMENT",
        "To deliver a resilient, ultra-low-latency (<1ms), and resource-efficient (<1MB) edge intrusion detection platform tailored for mining Operational Technology (OT) and SCADA environments in remote African and Russian resource facilities, safeguarding critical mineral assets and underground human lives against cyber-physical disruption."
    )

    # Section 1: Executive Summary & Document Control
    add_heading_1(doc, "1. EXECUTIVE SUMMARY & DOCUMENT CONTROL")
    add_body(doc,
        "This Product Requirements Document (PRD) defines the functional, non-functional, interface, and operational specifications for the 'Securing the Digital Mine' intrusion detection platform. Modern mineral extraction operations are deploying Industrial IoT (IIoT) sensors, automated haulage, and SCADA-governed milling circuits. However, existing IT-centric intrusion detection systems are unviable in remote mining sites due to heavy computational overhead, reliance on high-bandwidth cloud connectivity, and lack of adaptation to industrial protocol traffic (Modbus, DNP3, OPC-UA)."
    )
    
    add_formatted_table(doc,
        ["Document Attribute", "Specification Details"],
        [
            ["Product Name", "Securing the Digital Mine - Edge IDS Platform"],
            ["Package Identifier", "@mhiskall282/unesco-mine-sec-cli (GitHub Packages)"],
            ["Target Hardware", "Raspberry Pi 4B/5 (1GB-4GB RAM), Industrial Gateways, AWS EC2"],
            ["Target Protocols", "Modbus RTU/TCP, DNP3, OPC-UA, Ethernet/IP, MQTT"],
            ["Lead Institution", "University of Education, Winneba (UEW) & UEW Innovation Hub"],
            ["Forum Submission", "Russian-African Forum-Contest of Young Scientists 2026 (UNESCO)"],
            ["Document Version", "v3.0.0 (Production Candidate)"],
            ["Release Date", "August 2026"]
        ],
        col_widths=[2.5, 4.0]
    )

    # Section 2: User Personas & Problem Scenarios
    add_heading_1(doc, "2. TARGET USER PERSONAS & PAIN POINTS")
    
    add_heading_2(doc, "Persona 1: Kwame Mensah - Edge OT / SCADA Maintenance Engineer")
    add_body(doc, "Role: Manages programmable logic controllers (PLCs), sensor loops, and telemetry hubs at a gold processing plant in Tarkwa, Ghana.")
    add_bullet(doc, "Pain Point: Remote extraction sites have intermittent satellite connections and solar/battery power limits. Cloud-based security tools fail during link dropouts.")
    add_bullet(doc, "Need: A plug-and-play CLI agent that runs locally on existing 1GB RAM Raspberry Pi gateways with zero cloud dependency and sub-millisecond local inference.")

    add_heading_2(doc, "Persona 2: Elena Petrova - Industrial Cybersecurity Operations Lead")
    add_body(doc, "Role: Oversees central security monitoring and incident triage across subsoil extraction complexes.")
    add_bullet(doc, "Pain Point: Traditional intrusion detection systems generate thousands of cryptic numerical alerts with high false alarm rates that overwhelm control room operators.")
    add_bullet(doc, "Need: Clear, human-readable attack labels (e.g., 'DoS', 'Probe', 'Modbus Injection') with calibrated confidence percentages and real-time dashboard visualization.")

    add_heading_2(doc, "Persona 3: Mine Operations Director / General Manager")
    add_body(doc, "Role: Accountable for production output, regulatory safety compliance, and operational capital expenditure.")
    add_bullet(doc, "Pain Point: Unplanned mill or ventilation downtime costs between USD $50,000 and $500,000 per hour. Proprietary enterprise IDS solutions require exorbitant annual licensing fees.")
    add_bullet(doc, "Need: An open-source, mathematically proven security architecture that delivers tangible return on investment (>200x) and meets Minerals Commission safety mandates.")

    # Section 3: Product Architecture & System Layers
    add_heading_1(doc, "3. SYSTEM ARCHITECTURE & CORE PIPELINE")
    add_image_figure(doc, "research/figures/system_architecture.png", "Figure 3.1: Four-Layer Edge-to-Cloud System Architecture", width_inches=6.2)
    add_body(doc,
        "The product architecture consists of four tightly integrated layers:\n"
        "1. Layer 1 (Industrial Ingestion Layer): Real-time network sniffer (@mhiskall282/unesco-mine-sec-cli) hooking promiscuous interfaces.\n"
        "2. Layer 2 (Metaheuristic Optimization Layer): Binary Whale Optimization Algorithm (BWOA) pruning 41 features down to 10 key features.\n"
        "3. Layer 3 (Deep Learning Classification Layer): Spatial-temporal Conv1D-LSTM neural classifier compressed via Float16 quantization to 0.82 MB.\n"
        "4. Layer 4 (Operational & Alerting Layer): Asynchronous FastAPI inference engine and multi-tenant Laravel Livewire SaaS dashboard."
    )

    # Section 4: Functional Requirements (FR)
    add_heading_1(doc, "4. FUNCTIONAL REQUIREMENTS SPECIFICATION")
    add_formatted_table(doc,
        ["Req ID", "Requirement Name", "Priority", "Detailed Functional Description"],
        [
            ["FR-01", "Promiscuous Packet Sniffing", "P0 (Must Have)", "The edge CLI agent must bind to any active network interface (Ethernet, Wi-Fi, virtual TAP) and capture raw bidirectional network frames up to 1,000 flows/sec."],
            ["FR-02", "BWOA Feature Pruning", "P0 (Must Have)", "The system must automatically isolate and extract exactly the 10 BWOA-selected features from raw packet headers, achieving a 75.61% telemetry data reduction."],
            ["FR-03", "Multi-Class Attack Classification", "P0 (Must Have)", "The deep learning inference engine must classify incoming flows into five distinct categories: Normal, Denial of Service (DoS), Network Probing, Remote-to-Local (R2L), and User-to-Root (U2R)."],
            ["FR-04", "Float16 Quantized Inference", "P0 (Must Have)", "The model must execute via a quantized Float16 TensorFlow Lite engine to achieve single-sample inference latency of less than 1.0 millisecond on ARM Cortex-A72 hardware."],
            ["FR-05", "Interactive CLI Wizard", "P1 (Should Have)", "When launched without parameters, the CLI sniffer must prompt the operator interactively for target API URL, network interface, and authentication bearer token."],
            ["FR-06", "Headless Daemon Execution", "P1 (Should Have)", "The CLI sniffer must support non-interactive execution via command-line flags (--url, --key, --interface) for integration into Linux systemd services and cron schedules."],
            ["FR-07", "Live Dashboard Alerting", "P1 (Should Have)", "The Laravel SaaS dashboard must ingest flow telemetry via /api/external/analyze and broadcast real-time Livewire events with confidence scores and threat severity levels."],
            ["FR-08", "Compliance Audit Trail", "P2 (Nice to Have)", "All detected anomalies, affected IP endpoints, and operator acknowledgment timestamps must be recorded in an immutable audit log for regulatory compliance reporting."]
        ],
        col_widths=[0.9, 1.8, 1.2, 2.6]
    )

    # Section 5: Non-Functional Requirements (NFR)
    add_heading_1(doc, "5. NON-FUNCTIONAL REQUIREMENTS SPECIFICATION")
    add_formatted_table(doc,
        ["Req ID", "Category", "Target Metric", "Verification Method"],
        [
            ["NFR-01", "Inference Latency", "< 1.0 ms on edge hardware (Raspberry Pi 4B)", "Benchmarked at 0.76 ms over 1,000 inference runs (PASS)"],
            ["NFR-02", "Memory Footprint", "< 512 MB peak RAM on 1GB RAM edge gateway", "Profiled at 290.31 MB peak memory consumption (PASS)"],
            ["NFR-03", "Model Storage Size", "< 1.0 MB compressed TFLite binary", "Measured at exactly 0.82 MB (83.2% compression ratio, PASS)"],
            ["NFR-04", "Power Efficiency", "< 3.0 Watts continuous power draw", "Measured at 2.5 Watts under continuous load, solar/battery ready (PASS)"],
            ["NFR-05", "Classification Accuracy", "> 70.0% multi-class accuracy on KDDTest+", "Achieved 70.56% accuracy and 0.7127 Macro F1 on 22,544 samples (PASS)"],
            ["NFR-06", "Benign Precision", "> 95.0% precision on normal mining traffic", "Achieved 96.89% precision, preventing false operational shutdowns (PASS)"],
            ["NFR-07", "DoS Attack Recall", "> 85.0% recall on volumetric attack streams", "Achieved 89.04% recall on DoS intrusions (PASS)"],
            ["NFR-08", "Cloud Scalability & Throughput", "> 500 req/s throughput on AWS EC2", "Benchmarked at 617.13 req/s, 1.57 ms mean latency on AWS EC2 (PASS)"]
        ],
        col_widths=[1.0, 1.4, 2.3, 1.8]
    )

    # Section 6: Release Roadmap & Success Metrics
    add_heading_1(doc, "6. PRODUCT RELEASE ROADMAP & SUCCESS METRICS")
    add_formatted_table(doc,
        ["Phase", "Milestone Name", "Timeline", "Core Deliverables & KPIs"],
        [
            ["Phase 1", "Algorithm Design & Benchmark Validation", "Months 1 - 3", "BWOA optimization, CNN-LSTM training, 75 unit test suite, Float16 quantization. Target: >70% accuracy, <1ms latency (COMPLETED)."],
            ["Phase 2", "Edge & Cloud Deployment & Empirical Benchmarking", "Months 4 - 6", "NPM package (@mhiskall282/unesco-mine-sec-cli), SWaT transfer learning, Raspberry Pi & AWS EC2 1-command deployers, live cloud empirical benchmarking (617 req/s, 1.57ms latency) and publication data export (COMPLETED)."],
            ["Phase 3", "On-Site Field Validation & Multi-Mine Rollout", "Months 7 - 12", "Live SCADA capture at Gold Fields Tarkwa, federated multi-concession learning, automated Minerals Commission ESG audit reporting."]
        ],
        col_widths=[1.0, 1.8, 1.2, 2.5]
    )

    output_path = "research/PRD.docx"
    doc.save(output_path)
    print(f"Product Requirements Document (PRD) saved successfully to {output_path}!")

if __name__ == "__main__":
    create_prd()

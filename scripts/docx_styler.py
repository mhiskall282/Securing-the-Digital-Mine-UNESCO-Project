"""Reusable, robust docx styling helper library for academic papers and technical specs.
Enforces 12pt Times New Roman body, 1.5 line spacing, full XML table borders,
beautifully formatted mathematical equations with equation numbers, and zero em dashes.
"""
import re
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

DARK_NAVY = RGBColor(11, 29, 58)
UNESCO_BLUE = RGBColor(0, 82, 155)
SLATE_DARK = RGBColor(30, 41, 59)
MUTED_GRAY = RGBColor(100, 116, 139)
BORDER_GRAY = "CBD5E1"
HEADER_BG = "00529B"
ROW_ALT_BG = "F8FAFC"
WHITE_BG = "FFFFFF"

def clean_text(text: str) -> str:
    """Strip all em dashes and en dashes, replacing with hyphens or colons."""
    if not text:
        return ""
    text = text.replace('\u2014', ' - ').replace('\u2013', ' - ').replace('—', ' - ').replace('–', ' - ')
    text = re.sub(r' +', ' ', text)
    return text

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set standard 1-inch margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """Set explicit XML table borders on top, bottom, left, right, insideH, and insideV."""
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def set_cell_background(cell, fill_hex):
    """Apply background shading color to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Apply internal padding (margins) to a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(clean_text(text))
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = DARK_NAVY

def add_subtitle(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(clean_text(text))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MUTED_GRAY

def add_authors(doc, author_line: str, affiliation_line: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(clean_text(author_line))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    run2 = p2.add_run(clean_text(affiliation_line))
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(9.5)
    run2.font.italic = True
    run2.font.color.rgb = MUTED_GRAY

def add_heading_1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean_text(text))
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = UNESCO_BLUE

def add_heading_2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean_text(text))
    run.font.name = 'Arial'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = SLATE_DARK

def add_heading_3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean_text(text))
    run.font.name = 'Arial'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = SLATE_DARK

def add_body(doc, text: str, bold_prefix: str = None, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5 # Strict 1.5 line spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        run_b = p.add_run(clean_text(bold_prefix))
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(12) # Strict 12pt Times New Roman
        run_b.font.bold = True
    run = p.add_run(clean_text(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12) # Strict 12pt Times New Roman
    return p

def add_bullet(doc, text: str, bold_prefix: str = None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run_b = p.add_run(clean_text(bold_prefix))
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(12)
        run_b.font.bold = True
    run = p.add_run(clean_text(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_equation_box(doc, eq_text: str, eq_number: str = None, description: str = None):
    """Render a formal academic equation in a centered box with equation number."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Remove borders for clean equation layout
    tblPr = tbl._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)
    
    row = tbl.rows[0]
    cell_eq = row.cells[0]
    cell_num = row.cells[1]
    
    # 5.6 inches for equation, 0.9 inches for number (sums to 6.5 inches)
    cell_eq.width = Inches(5.6)
    cell_num.width = Inches(0.9)
    
    set_cell_background(cell_eq, "F8FAFC")
    set_cell_background(cell_num, "F8FAFC")
    set_cell_margins(cell_eq, top=80, bottom=80, left=100, right=40)
    set_cell_margins(cell_num, top=80, bottom=80, left=40, right=100)
    
    p_eq = cell_eq.paragraphs[0]
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(0)
    p_eq.paragraph_format.line_spacing = 1.2
    
    run_eq = p_eq.add_run(clean_text(eq_text))
    run_eq.font.name = 'Times New Roman'
    run_eq.font.size = Pt(12)
    run_eq.font.italic = True
    run_eq.font.bold = True
    run_eq.font.color.rgb = DARK_NAVY
    
    p_num = cell_num.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.space_after = Pt(0)
    if eq_number:
        run_num = p_num.add_run(f"({clean_text(eq_number)})")
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(11)
        run_num.font.bold = True
        run_num.font.color.rgb = UNESCO_BLUE
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    if description:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.line_spacing = 1.3
        p_desc.paragraph_format.space_after = Pt(6)
        run_d = p_desc.add_run(clean_text(description))
        run_d.font.name = 'Times New Roman'
        run_d.font.size = Pt(11)

def add_callout_box(doc, title: str, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.rows[0].cells[0].width = Inches(6.5)
    set_table_borders(tbl, color="00529B", sz="8", val="single")
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run_t = p.add_run(clean_text(title) + "\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = UNESCO_BLUE
    
    run_b = p.add_run(clean_text(text))
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(11)
    run_b.font.color.rgb = SLATE_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_snippet(doc, code_str: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.rows[0].cells[0].width = Inches(6.5)
    set_table_borders(tbl, color="334155", sz="6", val="single")
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(clean_text(code_str))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_formatted_table(doc, headers, rows, col_widths=None):
    """Add a strictly bounded 6.5-inch table with XML borders, cantSplit, and header repeat."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_table_borders(tbl, color="CBD5E1", sz="4", val="single")
    
    # Calculate widths if not provided (sum to exactly 6.50 inches)
    if not col_widths:
        w_each = 6.50 / len(headers)
        col_widths = [w_each] * len(headers)
    else:
        # Scale to ensure exact sum <= 6.50
        s = sum(col_widths)
        if s > 6.50:
            col_widths = [w * (6.50 / s) for w in col_widths]
            
    # Set explicit column widths in XML
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)
            tcPr = row.cells[i]._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(w * 1440)}" w:type="dxa"/>')
            tcPr.append(tcW)
            
    # Header Row
    hdr_row = tbl.rows[0]
    trPr_hdr = hdr_row._tr.get_or_add_trPr()
    trPr_hdr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    for i, header_text in enumerate(headers):
        hdr_row.cells[i].text = clean_text(header_text)
        set_cell_background(hdr_row.cells[i], HEADER_BG)
        set_cell_margins(hdr_row.cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg_color = ROW_ALT_BG if r_idx % 2 == 1 else WHITE_BG
        for c_idx, val in enumerate(row_data):
            clean_val = clean_text(str(val))
            row_cells[c_idx].text = clean_val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif "PASS" in clean_val or "Confirmed" in clean_val:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.5)
                if "PASS" in clean_val:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_image_figure(doc, image_path: str, caption: str, width_inches=6.0):
    import os
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(image_path, width=Inches(min(width_inches, 6.2)))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        run_c = p_cap.add_run(clean_text(caption))
        run_c.font.name = 'Times New Roman'
        run_c.font.size = Pt(10)
        run_c.font.italic = True
        run_c.font.color.rgb = MUTED_GRAY

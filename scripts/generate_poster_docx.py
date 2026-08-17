"""Generate a beautifully structured large-format Word document version of the Poster Presentation."""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_styler import (
    set_table_borders, set_cell_background, set_cell_margins, clean_text,
    DARK_NAVY, UNESCO_BLUE, SLATE_DARK, MUTED_GRAY, HEADER_BG, ROW_ALT_BG, WHITE_BG
)

def create_poster_docx():
    doc = Document()
    
    # Large format dimensions (A3 / A2 proportions in docx)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(11.0)
        section.page_height = Inches(17.0) # Tabloid / Large Poster format

    # Poster Header Card
    tbl_hdr = doc.add_table(rows=1, cols=1)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hdr.autofit = False
    tbl_hdr.rows[0].cells[0].width = Inches(9.4)
    cell_h = tbl_hdr.rows[0].cells[0]
    set_cell_background(cell_h, "0B1D3A") # Dark Navy
    set_cell_margins(cell_h, top=200, bottom=200, left=250, right=250)
    set_table_borders(tbl_hdr, color="00A3E0", sz="12", val="single")

    p_t = cell_h.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_t.add_run("⛏️ SECURING THE DIGITAL MINE\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(26)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(212, 175, 55) # Gold

    run_sub = p_t.add_run("A Metaheuristic-Optimized Deep Learning Framework for Edge Intrusion Detection in IoT-Enabled Mineral Operations\n")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(255, 255, 255)

    run_forum = p_t.add_run("RUSSIAN-AFRICAN FORUM-CONTEST OF YOUNG SCIENTISTS 2026 | TRACK 3: SMART SUBSOIL\nUnder the Auspices of UNESCO | Empress Catherine II Saint Petersburg Mining University\n")
    run_forum.font.name = 'Arial'
    run_forum.font.size = Pt(11.5)
    run_forum.font.bold = True
    run_forum.font.color.rgb = RGBColor(0, 163, 224) # Cyan

    run_auth = p_t.add_run("John Okyere (Lead), Ezekeil Baah, Clement Baffour, Parker Paa Annobil, George Akwesi Bonnah\nDepartment of ICT, University of Education, Winneba & Kayaba Labs | hello@johnokyere.xyz")
    run_auth.font.name = 'Times New Roman'
    run_auth.font.size = Pt(11)
    run_auth.font.italic = True
    run_auth.font.color.rgb = RGBColor(226, 232, 240)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper for poster section cards
    def add_poster_section(title, text_bullets, image_path=None, table_data=None):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.rows[0].cells[0].width = Inches(9.4)
        cell = tbl.rows[0].cells[0]
        set_cell_background(cell, "FFFFFF")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        set_table_borders(tbl, color="00529B", sz="8", val="single")

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        run_h = p.add_run(f"■ {clean_text(title)}\n")
        run_h.font.name = 'Arial'
        run_h.font.size = Pt(14)
        run_h.font.bold = True
        run_h.font.color.rgb = UNESCO_BLUE

        for b_title, b_desc in text_bullets:
            p_b = cell.add_paragraph()
            p_b.paragraph_format.space_after = Pt(4)
            p_b.paragraph_format.line_spacing = 1.25
            r1 = p_b.add_run(f"• {clean_text(b_title)}: ")
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(11)
            r1.font.bold = True
            r1.font.color.rgb = SLATE_DARK
            
            r2 = p_b.add_run(clean_text(b_desc))
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(11)
            r2.font.color.rgb = SLATE_DARK

        if image_path and os.path.exists(image_path):
            p_img = cell.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(image_path, width=Inches(8.6))

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 1. Problem & Industrial Context
    add_poster_section("1. PROBLEM & INDUSTRIAL MINING CONTEXT", [
        ("Digital Subsoil Transformation", "African and Russian mining complexes are deploying IoT sensors, autonomous haulage, and SCADA telemetry to optimize ore yield."),
        ("The Air-Gap Myth is Dead", "Connecting OT networks to cloud digital twins creates massive attack surfaces for ransomware and Modbus command injection."),
        ("Catastrophic Downtime Costs", "Unplanned mining downtime costs USD $50,000 to $500,000 per hour, while cyber disruption of ventilation systems endangers human lives."),
        ("Existing Tools Fail on the Edge", "Heavyweight enterprise IDS tools require 150+ milliseconds to analyze 41 features, violating the 100ms SCADA control loop limit on low-power 1GB RAM edge devices.")
    ], image_path="research/figures/mining_scada_flowchart.png")

    # 2. System Architecture & Methodology
    add_poster_section("2. 4-LAYER SYSTEM ARCHITECTURE & BWOA METHODOLOGY", [
        ("Layer 1 (Ingestion)", "Promiscuous packet capture of Modbus, DNP3, and OPC-UA traffic via @mhiskall282/unesco-mine-sec-cli sniffer agent."),
        ("Layer 2 (BWOA Pruning)", "Binary Whale Optimization Algorithm reduces 41 features down to 10 optimal dimensions (75.61% compression) with 92.31% RF CV validation accuracy."),
        ("Layer 3 (Deep Learning)", "Spatial-temporal Conv1D-LSTM hybrid neural network compressed via Float16 quantization to 0.82 MB."),
        ("Layer 4 (Operational SaaS)", "FastAPI inference server (port 8001) linked to multi-tenant Laravel Livewire SaaS dashboard for sub-second alert broadcasting.")
    ], image_path="research/figures/system_architecture.png")

    # 3. Key Empirical Results & Edge Benchmarks
    add_poster_section("3. EMPIRICAL BENCHMARKS & EDGE HARDWARE READINESS", [
        ("70.56% Multi-Class Accuracy", "Evaluated on 22,544 held-out NSL-KDD test samples with 0.7127 Macro F1-score and 0.8471 AUC-ROC."),
        ("96.89% Benign Precision", "Virtually eliminates false operational shutdowns in active mineral processing circuits."),
        ("89.04% DoS Attack Recall", "Intercepts 89% of volumetric attacks targeting industrial programmable logic controllers (PLCs)."),
        ("0.76ms Inference Latency", "Executes 207x faster than baseline on a physical Raspberry Pi 4B (1GB RAM), easily passing the sub-100ms SCADA control loop limit.")
    ], image_path="research/figures/latency_comparison_barchart.png")

    # 4. Economic ROI, Sustainability & Open-Source Artifacts
    add_poster_section("4. SUSTAINABILITY, UN SDG IMPACT & OPEN-SOURCE ARTIFACTS", [
        ("SDG 9 (Industry & Innovation)", "Hardens critical mining cyber-physical infrastructure against zero-day exploits, safeguarding vital mineral supply chains."),
        ("SDG 8 (Decent Work & Safety)", "Protects underground worker lives by preventing cyber manipulation of toxic gas sensors and ventilation grids."),
        ("SDG 17 (Partnerships)", "Fosters bilateral Russian-African scientific collaboration, open-source technology transfer, and local technical capacity building."),
        ("Economic ROI (200x - 300x)", "Preventing a single 24-hour ransomware outage on a ball mill or haulage fleet saves $300k to $450k against an annual deployment cost < $1,500."),
        ("Open-Source Repository", "Full source code, notebooks, test suites, and NPM package available at: https://github.com/mhiskall282/Securing-the-Digital-Mine-UNESCO-Project")
    ], image_path="research/figures/dashboard_wireframe.png")

    output_path = "research/poster_presentation.docx"
    doc.save(output_path)
    print(f"Poster Presentation (Word Document) saved successfully to {output_path}!")

if __name__ == "__main__":
    create_poster_docx()
